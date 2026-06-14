import os
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# ---------------------------------------------------------------------------
# XML HELPERS FOR ADVANCED WORD FORMATTING
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """Applies a background color to an individual table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    tcPr.append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets padding inside table cells (measured in dxa: 20 dxa = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar_xml = (
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(parse_xml(tcMar_xml))

def apply_table_styles(table):
    """Applies clean, horizontal-only borders (no vertical lines) for a modern look."""
    tblPr = table._tbl.tblPr
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="8" w:space="0" w:color="5C768D"/>'
        f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="5C768D"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))

def set_column_widths(table, widths):
    """Ensures each column in a table has a fixed width to prevent squishing."""
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = w

def format_cell(cell, text, font_size=9.5, bold=False, color_rgb=None, shading_hex=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """A comprehensive helper to format text, alignment, margins, and shading in a cell."""
    cell.text = ""  # Clear any default text
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(font_size)
    run.font.bold = bold
    
    if color_rgb:
        run.font.color.rgb = color_rgb
    else:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
    if shading_hex:
        set_cell_shading(cell, shading_hex)
        
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)

# ---------------------------------------------------------------------------
# DATA PROCESSING & CALCULATIONS
# ---------------------------------------------------------------------------

def load_and_process_data(excel_path):
    """Loads all 5 sheets and performs status and KPI calculations."""
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"Input file not found: {excel_path}")
        
    print(f"Loading data from: {excel_path}...")
    excel_file = pd.ExcelFile(excel_path)
    
    # Load sheets
    df_projects = pd.read_excel(excel_file, sheet_name="Projects")
    df_achievements = pd.read_excel(excel_file, sheet_name="Achievements")
    df_risks = pd.read_excel(excel_file, sheet_name="Risks")
    df_nextsteps = pd.read_excel(excel_file, sheet_name="NextSteps")
    df_decisions = pd.read_excel(excel_file, sheet_name="Decisions")
    
    # Calculate Project Status
    # Green: Delay <= 2 and Risk Score <= 2
    # Red: Delay > 5 or Risk Score > 3
    # Yellow: Otherwise
    def calc_status(row):
        delay = row["Delay Days"]
        risk = row["Risk Score"]
        if delay <= 2 and risk <= 2:
            return "Green"
        elif delay > 5 or risk > 3:
            return "Red"
        else:
            return "Yellow"
            
    df_projects["Status"] = df_projects.apply(calc_status, axis=1)
    
    # Generate One-Line Executive Summary Narratives
    def generate_narrative(row):
        proj = row["Project"]
        if proj == "CRM Upgrade":
            return "Progress slightly behind plan due to testing delays."
        elif proj == "ERP Rollout":
            return "On schedule and within budget."
        elif proj == "Data Migration":
            return "Critical schedule slippage impacting deployment."
        else:
            status = row["Status"]
            if status == "Green":
                return "On track with key milestones met and financials within variance guidelines."
            elif status == "Yellow":
                return f"Experiencing minor setbacks (Delay: {row['Delay Days']} days) with remediations underway."
            else:
                return f"Underperforming schedule and risk thresholds. Executive intervention recommended."

    df_projects["OneLineSummary"] = df_projects.apply(generate_narrative, axis=1)
    
    # Calculate KPIs
    # CPI = Budget / Actual Cost
    # SPI = Actual % / Planned %
    # Budget Consumed % = Actual Cost / Budget
    # Progress % = Actual % / Planned %
    df_projects["CPI"] = df_projects["Budget"] / df_projects["Actual Cost"]
    df_projects["SPI"] = df_projects["Actual %"] / df_projects["Planned %"]
    df_projects["Budget Consumed %"] = df_projects["Actual Cost"] / df_projects["Budget"]
    df_projects["Progress %"] = df_projects["Actual %"] / df_projects["Planned %"]
    
    return df_projects, df_achievements, df_risks, df_nextsteps, df_decisions

# ---------------------------------------------------------------------------
# REPORT GENERATION
# ---------------------------------------------------------------------------

def generate_word_report(df_projects, df_achievements, df_risks, df_nextsteps, df_decisions, output_path):
    """Generates the PMO Executive Report Word Document."""
    doc = Document()
    
    # Page Setup (1-inch margins)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Typography Setup (Normal Style)
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Helper for Headings with spacing
    def add_section_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        h.paragraph_format.keep_with_next = True
        if level == 1:
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
            h.style.font.size = Pt(14)
            h.style.font.bold = True
            h.style.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)  # Deep Navy
        elif level == 2:
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            h.style.font.size = Pt(11.5)
            h.style.font.bold = True
            h.style.font.color.rgb = RGBColor(0x5C, 0x76, 0x8D)  # Slate Gray
        return h

    # -----------------------------------------------------------------------
    # HEADER / TITLE BLOCK
    # -----------------------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("PMO WEEKLY EXECUTIVE REPORT")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(20)
    run_sub = p_sub.add_run(f"Generated on {datetime.now().strftime('%d-%b-%Y')} | Portfolio Status & Escalation Summary")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(10)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x5C, 0x76, 0x8D)
    
    # -----------------------------------------------------------------------
    # SECTION 1: OVERALL STATUS
    # -----------------------------------------------------------------------
    add_section_heading("1. Overall Status")
    
    for _, row in df_projects.iterrows():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        
        # Project Name
        run_proj = p.add_run(f"{row['Project']} – ")
        run_proj.bold = True
        
        # Status Label with Colors
        status = row['Status']
        run_status = p.add_run(status)
        run_status.bold = True
        if status == "Green":
            run_status.font.color.rgb = RGBColor(0x38, 0x57, 0x23)  # Green
        elif status == "Yellow":
            run_status.font.color.rgb = RGBColor(0xB9, 0x77, 0x0E)  # Dark Yellow
        else:
            run_status.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)  # Red
            
        # Optional metadata info
        run_meta = p.add_run(f" (Delay: {row['Delay Days']}d, Risk Score: {row['Risk Score']}/5)")
        run_meta.font.size = Pt(9)
        run_meta.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
        
        # One-line Summary
        p_sum = doc.add_paragraph()
        p_sum.paragraph_format.left_indent = Inches(0.2)
        p_sum.paragraph_format.space_before = Pt(0)
        p_sum.paragraph_format.space_after = Pt(8)
        run_sum = p_sum.add_run(row['OneLineSummary'])
        run_sum.italic = True

    # -----------------------------------------------------------------------
    # SECTION 2: WEEK ACHIEVEMENTS
    # -----------------------------------------------------------------------
    add_section_heading("2. Week Achievements")
    
    for _, proj_row in df_projects.iterrows():
        proj_name = proj_row['Project']
        p_proj = doc.add_paragraph()
        p_proj.paragraph_format.space_before = Pt(6)
        p_proj.paragraph_format.space_after = Pt(2)
        run_proj = p_proj.add_run(proj_name)
        run_proj.bold = True
        
        # Load achievements for this project
        achievements = df_achievements[df_achievements['Project'] == proj_name]
        if not achievements.empty:
            for _, ach_row in achievements.iterrows():
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_before = Pt(0)
                p_bullet.paragraph_format.space_after = Pt(2)
                p_bullet.add_run(ach_row['Achievement'])
        else:
            p_none = doc.add_paragraph()
            p_none.paragraph_format.left_indent = Inches(0.2)
            p_none.paragraph_format.space_after = Pt(2)
            run_none = p_none.add_run("No achievements recorded for this reporting period.")
            run_none.font.italic = True

    # -----------------------------------------------------------------------
    # SECTION 3: ACTIVE RISKS
    # -----------------------------------------------------------------------
    add_section_heading("3. Active Risks")
    
    # Filter High/Critical risks (Probability or Impact is High/Critical)
    hi_crit_risks = df_risks[
        df_risks['Probability'].isin(['High', 'Critical']) | 
        df_risks['Impact'].isin(['High', 'Critical'])
    ]
    
    if not hi_crit_risks.empty:
        headers = ["Project", "Risk Description", "Probability", "Impact", "Mitigation Plan"]
        widths = [Inches(1.2), Inches(1.8), Inches(0.8), Inches(0.8), Inches(1.9)]
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = 1
        table.autofit = False
        table.allow_autofit = False
        
        # Header formatting
        for i, h_text in enumerate(headers):
            format_cell(table.rows[0].cells[i], h_text, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF), shading_hex="1B365D")
            
        for _, r_row in hi_crit_risks.iterrows():
            row_cells = table.add_row().cells
            format_cell(row_cells[0], r_row['Project'], bold=True)
            format_cell(row_cells[1], r_row['Risk Description'])
            
            # Probability cell with red highlight if High/Critical
            prob = r_row['Probability']
            p_shd = "FCE4D6" if prob in ["High", "Critical"] else None
            p_color = RGBColor(0xC0, 0x00, 0x00) if prob in ["High", "Critical"] else None
            format_cell(row_cells[2], prob, color_rgb=p_color, shading_hex=p_shd, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Impact cell with red highlight if High/Critical
            imp = r_row['Impact']
            i_shd = "FCE4D6" if imp in ["High", "Critical"] else None
            i_color = RGBColor(0xC0, 0x00, 0x00) if imp in ["High", "Critical"] else None
            format_cell(row_cells[3], imp, color_rgb=i_color, shading_hex=i_shd, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            format_cell(row_cells[4], r_row['Mitigation'])
            
        set_column_widths(table, widths)
        apply_table_styles(table)
    else:
        p_none = doc.add_paragraph()
        p_none.add_run("No active High or Critical risks recorded.").font.italic = True

    # -----------------------------------------------------------------------
    # SECTION 4: NEXT STEPS
    # -----------------------------------------------------------------------
    add_section_heading("4. Next Steps")
    
    if not df_nextsteps.empty:
        headers = ["Project", "Task", "Owner", "Deadline"]
        widths = [Inches(1.2), Inches(2.5), Inches(1.0), Inches(1.8)]
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = 1
        table.autofit = False
        table.allow_autofit = False
        
        for i, h_text in enumerate(headers):
            format_cell(table.rows[0].cells[i], h_text, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF), shading_hex="1B365D")
            
        for _, n_row in df_nextsteps.iterrows():
            row_cells = table.add_row().cells
            format_cell(row_cells[0], n_row['Project'], bold=True)
            format_cell(row_cells[1], n_row['Task'])
            format_cell(row_cells[2], n_row['Owner'])
            
            # Date Formatting
            deadline = n_row['Deadline']
            if pd.notnull(deadline):
                if isinstance(deadline, (pd.Timestamp, datetime)):
                    deadline_str = deadline.strftime("%d-%b-%Y")
                else:
                    try:
                        deadline_str = pd.to_datetime(deadline).strftime("%d-%b-%Y")
                    except:
                        deadline_str = str(deadline)
            else:
                deadline_str = ""
                
            format_cell(row_cells[3], deadline_str, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
        set_column_widths(table, widths)
        apply_table_styles(table)
    else:
        p_none = doc.add_paragraph()
        p_none.add_run("No next steps recorded.").font.italic = True

    # -----------------------------------------------------------------------
    # SECTION 5: KPI DASHBOARD
    # -----------------------------------------------------------------------
    add_section_heading("5. KPI Dashboard")
    
    headers = ["Project", "CPI", "SPI", "Budget Consumed %", "Progress %"]
    widths = [Inches(1.5), Inches(1.0), Inches(1.0), Inches(1.5), Inches(1.5)]
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = 1
    table.autofit = False
    table.allow_autofit = False
    
    # Header format
    for i, h_text in enumerate(headers):
        align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        format_cell(table.rows[0].cells[i], h_text, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF), shading_hex="1B365D", alignment=align)
        
    def get_kpi_alert_colors(val):
        """Green >= 0.9, Yellow 0.8 - 0.89, Red < 0.8"""
        if val >= 0.9:
            return "E2F0D9", RGBColor(0x38, 0x57, 0x23)  # Soft green, dark green text
        elif val >= 0.8:
            return "FFF2CC", RGBColor(0x80, 0x60, 0x00)  # Soft yellow, dark yellow text
        else:
            return "FCE4D6", RGBColor(0xC0, 0x00, 0x00)  # Soft red, dark red text
            
    for _, k_row in df_projects.iterrows():
        row_cells = table.add_row().cells
        
        # Project Name
        format_cell(row_cells[0], k_row['Project'], bold=True)
        
        # CPI with shading
        cpi = k_row['CPI']
        cpi_shd, cpi_txt = get_kpi_alert_colors(cpi)
        format_cell(row_cells[1], f"{cpi:.2f}", color_rgb=cpi_txt, shading_hex=cpi_shd, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # SPI with shading
        spi = k_row['SPI']
        spi_shd, spi_txt = get_kpi_alert_colors(spi)
        format_cell(row_cells[2], f"{spi:.2f}", color_rgb=spi_txt, shading_hex=spi_shd, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Budget Consumed %
        budget_pct = k_row['Budget Consumed %']
        format_cell(row_cells[3], f"{budget_pct:.1%}", alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Progress %
        prog_pct = k_row['Progress %']
        format_cell(row_cells[4], f"{prog_pct:.1%}", alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
    set_column_widths(table, widths)
    apply_table_styles(table)

    # -----------------------------------------------------------------------
    # SECTION 6: DECISIONS REQUIRED
    # -----------------------------------------------------------------------
    add_section_heading("6. Decisions Required")
    
    if not df_decisions.empty:
        headers = ["Project", "Decision Required", "Context", "Options", "PM Recommendation"]
        widths = [Inches(1.2), Inches(1.8), Inches(1.2), Inches(1.1), Inches(1.2)]
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = 1
        table.autofit = False
        table.allow_autofit = False
        
        for i, h_text in enumerate(headers):
            format_cell(table.rows[0].cells[i], h_text, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF), shading_hex="1B365D")
            
        for _, d_row in df_decisions.iterrows():
            row_cells = table.add_row().cells
            format_cell(row_cells[0], d_row['Project'], bold=True)
            format_cell(row_cells[1], d_row['Decision Required'])
            format_cell(row_cells[2], d_row['Context'])
            format_cell(row_cells[3], d_row['Options'])
            format_cell(row_cells[4], d_row['Recommendation'])
            
        set_column_widths(table, widths)
        apply_table_styles(table)
    else:
        p_none = doc.add_paragraph()
        p_none.add_run("No decisions or escalations required at this time.").font.italic = True
        
    # Save the document
    doc.save(output_path)
    print(f"Executive Report successfully generated at: {output_path}")

# ---------------------------------------------------------------------------
# MAIN EXECUTION ENTRYPOINT
# ---------------------------------------------------------------------------

def run_report_generation(excel_path="PMO_Weekly_Report_Data.xlsx", output_path="PMO_Executive_Report.docx"):
    """Exposes report generation as a function that can be imported and run dynamically."""
    p_df, a_df, r_df, n_df, d_df = load_and_process_data(excel_path)
    generate_word_report(p_df, a_df, r_df, n_df, d_df, output_path)

if __name__ == "__main__":
    try:
        run_report_generation()
        print("PMO Executive Summary Automation Completed Successfully!")
    except Exception as e:
        print(f"Error during report generation: {e}")
        raise e
